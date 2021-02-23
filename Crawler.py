'''
Created on Oct 20, 2020

@author: sony.lap
'''
import docx
import openpyxl
import tweepy
import json
import re
from docx.document import Document
from pymongo import MongoClient

arr = []
jsonArray = []
consumer_key = 'p'
#'KHr9akjPxdRmJKNAsz8apEhvt'
consumer_secret = 'B-qa2-031cdd-0-302d0214496be84732a01f690268d3b8eb72e5b8ccf94e2202150085913117f2e1a8531505ee8ccfc8e98df3cf1748'
#'dcyMcL2rSMCR5UEFqdAtqmYKZUukHtVzObRwJTw9s7YAXhjCRm'
access_token = 'pu1'
#'3381765072-Rs4z8OxQjxQKNjl7HFelxjda3BxzICoGuHl5DSz'
access_token_secret = 'B-qa2031cbe-0-302d021500890ef262296563accd1cb4aab790323d2fd570d30214510bcdacdaa4f03f59477eef13f2af5ad13e3044'
#'koWj4hlDhfGg5MIgqN3v4WXdntsuni6Ljs7fvcVeqpppC'

# def getDB():
#     client = MongoClient('localhost:27017')
#     db = client.myFirstDB
#     return db

# def mongo_add(decoded, x, t):
#     db = getDB()
#     task = {
#         'id' : decoded['id'],
#         'location' : x,
#         'tweet' : t
#     }
#     db.tweet.insert_one({'location' : x})
#     k = db.tweet.find()
#     print(k.next())

#Adding to Excel File
def addExcel(decoded, x, t):
    excel = openpyxl.load_workbook('F:\Twiiterdata/tweet.xlsx')
    sheet1 = excel['Tweets']
    cell1 = sheet1.cell(row=1, column=1)
    if cell1.value is None:
        cell1.value = 'id'
        sheet1.cell(row=1, column=2).value = 'location'
        sheet1.cell(row=1, column=3).value = 'tweet'
    sheet1.cell(row=sheet1.max_row+1, column=1).value = decoded['id']                
    sheet1.cell(row=sheet1.max_row, column=2).value = x                
    sheet1.cell(row=sheet1.max_row, column=3).value = t
    excel.save('F:\Twiiterdata/tweet.xlsx')

#Adding to JSON File
def addJSON(decoded, x, t):
    task = {
        'id' : decoded['id'],
        'location' : x,
        'tweet' : t 
    }
    f = open('F:\Twiiterdata/tweet.json','r+')
    jsonArray = json.load(f)
    jsonArray.append(task)
    f.close()
    f = open('F:\Twiiterdata/tweet.json','w+')                
    f.write(json.dumps(jsonArray, indent=4))
    f.close()        

#Adding to Word File
def addWord(decoded, x, t):
    doc = docx.Document()
    doc.add_heading('Twitter - Tweet your way', 0)
    doc.add_paragraph("")
    doc.add_paragraph("Thanks " )
    run = doc.paragraphs[2].add_run()
    run.text = str(decoded['id'])
    run.bold = True
    run = doc.paragraphs[2].add_run()
    run.text = ' for tweeting, your tweet has been recorded'
    doc.add_paragraph('Keep posting for more followers and views')
    doc.save("F:\\Twiiterdata\\" + str(decoded['id']) + '.docx')


class StdOutListener(tweepy.StreamListener):
    def addressStandardization(self, city):
        cityS = (city.sub("([^A-Za-z0-9 \s]+))")    

    def on_data(self, data):
        decoded = json.loads(data)
        x = decoded['user']['location']
        if x is not None:
            if ',' in x:
                arr = x.split(',')
                x = re.sub("[^A-Za-z \s]", "", arr[1])
                t =  decoded['text']
                t = ' '.join(re.sub("(@[A-Za-z0-9]+)|([^0-9A-Za-z \t])|(\w+:\/\/\S+)"," ",t).split())
                t.replace("RT ", "", 1)

                STORING IN PDF but not incorporated duplicate logic yet
                addPdf(decoded, x, t)

                STORING IN WORD but not incorporated duplicate logic yet
                addWord(decoded, x, t)

                STORING IN EXCEL FILE
                addExcel(decoded, x, t)

                STORING IN JSON FILE
                addJSON(decoded, x, t)

                #STORING IN MONGO DB
                #mongo_add(decoded, x, t)

#         return True


#     def on_error(self, status_code):
#         if status_code == 420:
#             print('error connecting, try again later...')
#             return False

if __name__ == '__main__':
    auth = tweepy.OAuthHandler(consumer_key, consumer_secret)
    auth.set_access_token(access_token, access_token_secret)
    print("Showing all new tweets for #bollywood:")
    stream = tweepy.Stream(auth, StdOutListener())
    stream.filter(track=['bollywood'], is_async=True)
#     obj = StdOutListener()
#     obj.addressStandardization("Vigor!@31-=!`dfkdf123")